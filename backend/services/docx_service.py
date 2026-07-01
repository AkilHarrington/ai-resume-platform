# =========================================================
# File: docx_service.py
# Purpose: Generate .docx resume files via docxtpl (Jinja2 template injection)
#
# Architecture
# ────────────
#  1. _build_base_template()  — builds a Jinja2-tagged .docx template once per
#                               template name; visual styling is baked in here.
#  2. _TEMPLATE_CACHE         — module-level dict caches template bytes so
#                               each template is built at most once per process.
#  3. _build_context()        — converts parsed resume data into a Jinja2
#                               context dict (incl. RichText styled objects).
#  4. generate_resume_docx()  — loads cached template, renders with docxtpl,
#                               returns raw bytes.
#
# Templates
# ─────────
#  professional  — navy full-width banner headers, ATS-safe single column
#  modern        — blue accent headings with bottom-rule underlines
#  executive     — dark-green full-bleed shaded header + green banners
# =========================================================

import io
import re
import logging

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxtpl import DocxTemplate, RichText

from services.resume_parser import parse_resume_text

logger = logging.getLogger("ai_resume_studio")

_DATE_YEAR_RE = re.compile(r"\b(19|20)\d{2}\b")


# ─── Palette definitions ──────────────────────────────────────────────────────

_PALETTES: dict = {
    "professional": {
        "name_color":        RGBColor(0x1A, 0x35, 0x6E),   # navy
        "banner_bg":         "1A356E",
        "banner_text":       "FFFFFF",
        "title_color":       RGBColor(0x1A, 0x35, 0x6E),
        "company_color":     RGBColor(0x04, 0x78, 0x57),   # emerald
        "rule_color":        "047857",
        "date_color":        RGBColor(0x6B, 0x72, 0x80),
        "body_color":        RGBColor(0x1F, 0x1F, 0x1F),
        "section_style":     "banner",
        "name_align":        WD_ALIGN_PARAGRAPH.LEFT,
        "contact_align":     WD_ALIGN_PARAGRAPH.LEFT,
        "header_bg":         None,
        # hex strings for RichText.add(color=...)
        "title_color_hex":   "1A356E",
        "company_color_hex": "047857",
    },
    "modern": {
        "name_color":        RGBColor(0x00, 0x3D, 0x79),
        "banner_bg":         None,
        "banner_text":       "0078D7",
        "title_color":       RGBColor(0x00, 0x78, 0xD7),
        "company_color":     RGBColor(0x44, 0x44, 0x44),
        "rule_color":        "0078D7",
        "date_color":        RGBColor(0x6B, 0x72, 0x80),
        "body_color":        RGBColor(0x1F, 0x1F, 0x1F),
        "section_style":     "underline",
        "name_align":        WD_ALIGN_PARAGRAPH.LEFT,
        "contact_align":     WD_ALIGN_PARAGRAPH.LEFT,
        "header_bg":         None,
        "title_color_hex":   "0078D7",
        "company_color_hex": "444444",
    },
    "executive": {
        "name_color":        RGBColor(0xFF, 0xFF, 0xFF),    # white on dark bg
        "banner_bg":         "1C3A2B",                      # forest green
        "banner_text":       "FFFFFF",
        "title_color":       RGBColor(0x1C, 0x3A, 0x2B),
        "company_color":     RGBColor(0x16, 0x65, 0x34),   # emerald
        "rule_color":        "166534",
        "date_color":        RGBColor(0x6B, 0x72, 0x80),
        "body_color":        RGBColor(0x1F, 0x1F, 0x1F),
        "section_style":     "banner",
        "name_align":        WD_ALIGN_PARAGRAPH.CENTER,
        "contact_align":     WD_ALIGN_PARAGRAPH.CENTER,
        "header_bg":         "1C3A2B",
        "title_color_hex":   "1C3A2B",
        "company_color_hex": "166534",
    },
}


# ─── Template cache (in-memory; each template built once per process) ─────────

_TEMPLATE_CACHE: dict = {}


# ─── Low-level XML helpers ────────────────────────────────────────────────────

def _hex_to_rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _shade_paragraph(p, fill_hex: str) -> None:
    """Apply full-width background shading to a paragraph."""
    pPr = p._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:shd")):
        pPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.lstrip("#"))
    pPr.append(shd)


def _add_paragraph_border(p, side: str, color_hex: str,
                           sz: str = "6", space: str = "4") -> None:
    """Add a single-line border to one side of a paragraph."""
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), sz)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color_hex.lstrip("#"))
    pBdr.append(border)


def _set_right_tab(p, position: int = 9072) -> None:
    """Set a right-aligned tab stop (position in twips; default ≈ 6.3")."""
    pPr = p._p.get_or_add_pPr()
    tabs_el = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(position))
    tabs_el.append(tab)
    pPr.append(tabs_el)


# ─── Template paragraph helpers ───────────────────────────────────────────────

def _jinja_p(doc: Document, tag: str):
    """
    Add a marker paragraph containing ONLY a docxtpl paragraph-level control
    tag (e.g. '{%p if has_summary %}').  The run is 1pt white — invisible when
    the template file is opened in Word directly.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(tag)
    r.font.size      = Pt(1)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return p


def _banner_p(doc: Document, text: str, bg_hex: str, text_hex: str = "FFFFFF"):
    """Full-width shaded banner paragraph (Professional / Executive sections)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    _shade_paragraph(p, bg_hex)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "80")
    pPr.append(ind)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size      = Pt(8.5)
    r.font.color.rgb = _hex_to_rgb(text_hex)
    return p


def _underline_header_p(doc: Document, text: str, color_hex: str):
    """Colored uppercase label with a bottom-rule (Modern section style)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    _add_paragraph_border(p, "bottom", color_hex, sz="8", space="2")
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size      = Pt(9)
    r.font.color.rgb = _hex_to_rgb(color_hex)
    return p


def _section_header_p(doc: Document, text: str, palette: dict):
    if palette["section_style"] == "banner":
        return _banner_p(doc, text, palette["banner_bg"], palette["banner_text"])
    return _underline_header_p(doc, text, palette["banner_text"])


# ─── Base template builder ────────────────────────────────────────────────────

def _build_base_template(template_name: str) -> bytes:
    """
    Build a styled DOCX file with Jinja2 placeholder tags.

    Visual elements (shading, fonts, colors, borders) are baked into the
    paragraph/run formatting here.  Content is injected at render time by
    docxtpl using the context dict produced by _build_context().

    The file is serialised to bytes and cached — it is never written to disk.
    """
    pal = _PALETTES[template_name]
    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Inches(0.65)
        sec.bottom_margin = Inches(0.65)
        sec.left_margin   = Inches(0.80)
        sec.right_margin  = Inches(0.80)

    # Default body style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(2)

    # ── Name / contact header ─────────────────────────────────────────────────
    if pal["header_bg"]:
        # Executive: full-bleed dark header (name + headline + contact + stripe)
        p_name = doc.add_paragraph()
        p_name.alignment                = WD_ALIGN_PARAGRAPH.CENTER
        p_name.paragraph_format.space_before = Pt(12)
        p_name.paragraph_format.space_after  = Pt(2)
        _shade_paragraph(p_name, pal["header_bg"])
        r = p_name.add_run("{{ full_name }}")
        r.bold = True
        r.font.size      = Pt(22)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        _jinja_p(doc, "{%p if headline %}")
        p_hl = doc.add_paragraph()
        p_hl.alignment                = WD_ALIGN_PARAGRAPH.CENTER
        p_hl.paragraph_format.space_after = Pt(2)
        _shade_paragraph(p_hl, pal["header_bg"])
        r_hl = p_hl.add_run("{{ headline }}")
        r_hl.font.size      = Pt(11)
        r_hl.font.color.rgb = _hex_to_rgb("A7C4A0")  # sage on dark bg
        _jinja_p(doc, "{%p endif %}")

        _jinja_p(doc, "{%p if contact_line %}")
        p_c = doc.add_paragraph()
        p_c.alignment                 = WD_ALIGN_PARAGRAPH.CENTER
        p_c.paragraph_format.space_after = Pt(12)
        _shade_paragraph(p_c, pal["header_bg"])
        r_c = p_c.add_run("{{ contact_line }}")
        r_c.font.size      = Pt(9)
        r_c.font.color.rgb = _hex_to_rgb("A0AEC0")  # slate on dark bg
        _jinja_p(doc, "{%p endif %}")

        # Thin colour stripe below the header block
        stripe = doc.add_paragraph()
        stripe.paragraph_format.space_after = Pt(0)
        _shade_paragraph(stripe, pal["rule_color"])
        stripe.add_run("").font.size = Pt(3)

    else:
        # Professional / Modern: name, optional headline, optional contact, rule
        p_name = doc.add_paragraph()
        p_name.alignment                = pal["name_align"]
        p_name.paragraph_format.space_after = Pt(2)
        r = p_name.add_run("{{ full_name }}")
        r.bold = True
        r.font.size      = Pt(22)
        r.font.color.rgb = pal["name_color"]

        _jinja_p(doc, "{%p if headline %}")
        p_hl = doc.add_paragraph()
        p_hl.alignment                = pal["name_align"]
        p_hl.paragraph_format.space_after = Pt(2)
        r_hl = p_hl.add_run("{{ headline }}")
        r_hl.font.size      = Pt(11)
        r_hl.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        _jinja_p(doc, "{%p endif %}")

        _jinja_p(doc, "{%p if contact_line %}")
        p_c = doc.add_paragraph()
        p_c.alignment                 = pal.get("contact_align", WD_ALIGN_PARAGRAPH.LEFT)
        p_c.paragraph_format.space_after = Pt(4)
        r_c = p_c.add_run("{{ contact_line }}")
        r_c.font.size      = Pt(9)
        r_c.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        _jinja_p(doc, "{%p endif %}")

        rule_p = doc.add_paragraph()
        rule_p.paragraph_format.space_after = Pt(2)
        _add_paragraph_border(rule_p, "bottom", pal["rule_color"], sz="8", space="1")

    # ── Professional Summary ───────────────────────────────────────────────────
    _jinja_p(doc, "{%p if has_summary %}")
    _section_header_p(doc, "Professional Summary", pal)
    p_sum = doc.add_paragraph()
    p_sum.paragraph_format.space_after = Pt(4)
    r_sum = p_sum.add_run("{{ summary }}")
    r_sum.font.size      = Pt(10)
    r_sum.font.color.rgb = pal["body_color"]
    _jinja_p(doc, "{%p endif %}")

    # ── Experience ────────────────────────────────────────────────────────────
    _jinja_p(doc, "{%p if has_experience %}")
    _section_header_p(doc, "Experience", pal)
    _jinja_p(doc, "{%p for job in experience %}")

    # Job header: RichText label (company | title) + right-tab + date
    p_job = doc.add_paragraph()
    p_job.paragraph_format.space_before = Pt(6)
    p_job.paragraph_format.space_after  = Pt(0)
    _set_right_tab(p_job, position=9072)

    # label_rt is a RichText object injected by _build_context() — docxtpl
    # replaces the run with the multi-run formatted output automatically.
    r_label = p_job.add_run("{{ job.label_rt }}")
    r_label.font.size = Pt(10)

    # Tab character + date string (string substitution; run formatting is preserved)
    r_date = p_job.add_run("\t{{ job.date_range }}")
    r_date.font.size      = Pt(9)
    r_date.font.color.rgb = pal["date_color"]
    r_date.italic         = True

    # Bullet points (inner loop)
    _jinja_p(doc, "{%p for bullet in job.bullets %}")
    p_bul = doc.add_paragraph()
    p_bul.paragraph_format.space_after = Pt(1)
    p_bul.paragraph_format.left_indent = Inches(0.18)
    r_bul = p_bul.add_run("• {{ bullet }}")
    r_bul.font.size      = Pt(10)
    r_bul.font.color.rgb = pal["body_color"]
    _jinja_p(doc, "{%p endfor %}")  # end bullet loop

    _jinja_p(doc, "{%p endfor %}")  # end job loop
    _jinja_p(doc, "{%p endif %}")   # end has_experience

    # ── Education ─────────────────────────────────────────────────────────────
    _jinja_p(doc, "{%p if has_education %}")
    _section_header_p(doc, "Education", pal)
    _jinja_p(doc, "{%p for edu in education %}")

    p_inst = doc.add_paragraph()
    p_inst.paragraph_format.space_before = Pt(4)
    p_inst.paragraph_format.space_after  = Pt(1)
    r_inst = p_inst.add_run("{{ edu.institution }}")
    r_inst.bold          = True
    r_inst.font.size     = Pt(10)
    r_inst.font.color.rgb = pal["title_color"]

    _jinja_p(doc, "{%p if edu.degree_str %}")
    p_deg = doc.add_paragraph()
    p_deg.paragraph_format.space_after = Pt(2)
    r_deg = p_deg.add_run("{{ edu.degree_str }}")
    r_deg.font.size      = Pt(10)
    r_deg.font.color.rgb = pal["body_color"]
    _jinja_p(doc, "{%p endif %}")

    _jinja_p(doc, "{%p endfor %}")
    _jinja_p(doc, "{%p endif %}")

    # ── Skills ────────────────────────────────────────────────────────────────
    _jinja_p(doc, "{%p if has_skills %}")
    _section_header_p(doc, "Skills", pal)
    p_sk = doc.add_paragraph()
    p_sk.paragraph_format.space_after = Pt(4)
    r_sk = p_sk.add_run("{{ skills_line }}")
    r_sk.font.size      = Pt(10)
    r_sk.font.color.rgb = pal["body_color"]
    _jinja_p(doc, "{%p endif %}")

    # ── Certifications ────────────────────────────────────────────────────────
    _jinja_p(doc, "{%p if has_certifications %}")
    _section_header_p(doc, "Certifications", pal)
    _jinja_p(doc, "{%p for cert in certifications %}")
    p_cert = doc.add_paragraph()
    p_cert.paragraph_format.space_after = Pt(1)
    p_cert.paragraph_format.left_indent = Inches(0.18)
    r_cert = p_cert.add_run("• {{ cert.name }}")
    r_cert.font.size      = Pt(10)
    r_cert.font.color.rgb = pal["body_color"]
    _jinja_p(doc, "{%p endfor %}")
    _jinja_p(doc, "{%p endif %}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _get_template(template_name: str) -> bytes:
    """Return cached template bytes, building on first access."""
    if template_name not in _TEMPLATE_CACHE:
        _TEMPLATE_CACHE[template_name] = _build_base_template(template_name)
    return _TEMPLATE_CACHE[template_name]


# ─── Orphaned date-entry merge ────────────────────────────────────────────────

def _post_process_experience(experience: list) -> list:
    """
    Merge entries where the parser mis-identified a standalone date line
    (e.g. "March 2021 – Present") as a new company.  Their bullets belong to
    the preceding real experience entry.
    """
    if not experience:
        return experience
    merged: list = []
    for entry in experience:
        company = (entry.get("company") or "").strip()
        title   = (entry.get("title")   or "").strip()
        bullets = entry.get("bullets", [])
        is_date_entry = (
            not title
            and bool(_DATE_YEAR_RE.search(company))
            and any(
                tok in company
                for tok in ("–", "—", "-", "Present", "Current", "present", "current")
            )
        )
        if is_date_entry and merged:
            merged[-1]["bullets"] = merged[-1].get("bullets", []) + bullets
        else:
            merged.append(dict(entry))
    return merged


# ─── Context builder ──────────────────────────────────────────────────────────

def _build_context(resume_data: dict, template_name: str) -> dict:
    """
    Convert parsed resume data into a Jinja2 context dict for docxtpl.

    RichText objects are used for experience labels so company and job title
    can each carry their own colour, matching the template's palette.
    """
    pal = _PALETTES[template_name]

    full_name  = resume_data.get("fullName", "") or ""
    headline   = resume_data.get("headline", "") or ""
    contact    = resume_data.get("contact", {}) or {}
    summary    = resume_data.get("summary", "") or ""
    skills     = resume_data.get("skills", []) or []
    experience = _post_process_experience(resume_data.get("experience", []) or [])
    education  = resume_data.get("education", []) or []
    certs      = resume_data.get("certifications", []) or []

    # Single-line contact string
    contact_parts = [v for v in (
        contact.get("email"),
        contact.get("phone"),
        contact.get("location"),
        contact.get("linkedin"),
    ) if v]
    contact_line = "  |  ".join(contact_parts)

    # Experience — RichText label carries per-template colours
    exp_ctx: list = []
    for job in experience:
        company    = (job.get("company")   or "").strip()
        title      = (job.get("title")     or "").strip()
        start      = (job.get("startDate") or "").strip()
        end        = (job.get("endDate")   or "").strip()
        date_range = f"{start} – {end}" if (start and end) else (start or end)
        bullets    = [b.strip() for b in job.get("bullets", []) if b.strip()]

        if not company and not title and not bullets:
            continue  # skip entirely empty entries

        # Build RichText: company (accent colour) | title (primary colour)
        # size=20 = 10 pt in half-points (docxtpl RichText.add uses half-points)
        label_rt = RichText()
        if company:
            label_rt.add(company, bold=True, color=pal["company_color_hex"], size=20)
        if company and title:
            label_rt.add("  |  ", bold=False, color="888888", size=20)
        if title:
            label_rt.add(title, bold=True, color=pal["title_color_hex"], size=20)

        exp_ctx.append({
            "label_rt":   label_rt,
            "date_range": date_range,
            "bullets":    bullets,
        })

    # Education
    edu_ctx: list = []
    for edu in education:
        inst = (edu.get("institution")   or "").strip()
        deg  = (edu.get("degree")        or "").strip()
        gd   = (edu.get("graduationDate") or "").strip()
        if not inst:
            continue
        degree_str = f"{deg}  ({gd})" if (deg and gd) else deg
        edu_ctx.append({"institution": inst, "degree_str": degree_str})

    # Skills — join into a single bullet-separated line
    skills_line = "  •  ".join(str(s) for s in skills if s)

    # Certifications
    cert_ctx: list = []
    for c in certs:
        name = (c.get("name", "") if isinstance(c, dict) else str(c)).strip()
        if name:
            cert_ctx.append({"name": name})

    return {
        "full_name":          full_name,
        "headline":           headline,
        "contact_line":       contact_line,
        "summary":            summary,
        "experience":         exp_ctx,
        "education":          edu_ctx,
        "skills_line":        skills_line,
        "certifications":     cert_ctx,
        "has_summary":        bool(summary),
        "has_experience":     bool(exp_ctx),
        "has_education":      bool(edu_ctx),
        "has_skills":         bool(skills_line),
        "has_certifications": bool(cert_ctx),
    }


# ─── Fallback renderer (parser extracted nothing useful) ─────────────────────

def _fallback_raw_text(resume_text: str, template_name: str) -> bytes:
    """Line-by-line fallback when the structured parser yields no sections."""
    pal = _PALETTES[template_name]
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.65)
        sec.bottom_margin = Inches(0.65)
        sec.left_margin   = Inches(0.80)
        sec.right_margin  = Inches(0.80)

    KNOWN_HEADINGS = {
        "summary", "professional summary", "experience",
        "professional experience", "work experience", "education",
        "skills", "core competencies", "certifications", "licenses",
    }
    for line in resume_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        lower = stripped.lower()
        is_heading = (stripped.isupper() and 3 < len(stripped) < 60) or lower in KNOWN_HEADINGS
        if is_heading:
            if pal["section_style"] == "banner":
                _banner_p(doc, stripped, pal["banner_bg"], pal["banner_text"])
            else:
                _underline_header_p(doc, stripped, pal["banner_text"])
        elif stripped[0] in ("•", "-", "–", "*"):
            pp = doc.add_paragraph()
            pp.paragraph_format.left_indent = Inches(0.18)
            pp.add_run(stripped).font.size = Pt(10)
        else:
            pp = doc.add_paragraph()
            pp.add_run(stripped).font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── Public entry point ───────────────────────────────────────────────────────

def generate_resume_docx(resume_text: str, template: str = "professional") -> bytes:
    """
    Parse resume_text and render a styled .docx file.

    Uses docxtpl (Jinja2 template injection): visual design is defined once in
    a programmatically generated base template; each call injects resume data
    into that template via a Jinja2 context dict.

    Args:
        resume_text: Raw resume plain text (as extracted from upload).
        template:    "professional" | "modern" | "executive"

    Returns:
        Raw bytes of the rendered .docx file.
    """
    template = template.lower().strip()
    if template not in _PALETTES:
        template = "professional"

    resume_data = parse_resume_text(resume_text)
    context     = _build_context(resume_data, template)

    has_content = (
        context["has_summary"]
        or context["has_experience"]
        or context["has_education"]
        or context["has_skills"]
        or context["has_certifications"]
    )

    if not has_content:
        logger.warning(
            "generate_resume_docx: parser extracted no sections — "
            "fallback raw-text renderer (template=%s)", template
        )
        return _fallback_raw_text(resume_text, template)

    # Get (or lazily build) the cached base template, then render with docxtpl
    tpl_bytes = _get_template(template)
    tpl = DocxTemplate(io.BytesIO(tpl_bytes))
    tpl.render(context)

    out = io.BytesIO()
    tpl.save(out)
    result = out.getvalue()

    logger.info(
        "generate_resume_docx: %d bytes (template=%s)", len(result), template
    )
    return result
