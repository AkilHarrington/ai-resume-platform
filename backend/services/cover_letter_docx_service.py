# =========================================================
# File: services/cover_letter_docx_service.py
# Purpose:
# Generate cover letter DOCX files from plain-text cover letter output.
#
# Responsibilities:
# - format the generated cover letter into a Word document
# - apply basic typography and paragraph styling
# - save generated DOCX files for frontend download
#
# Key Notes:
# - this file handles cover letter export formatting only
# - it should not contain AI generation logic
# - the cover letter text is generated upstream by cover_letter_service.py
# - frontend downloads the generated file through the backend route
# =========================================================

# =========================================================
# Imports
# =========================================================

import os
import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


# =========================================================
# Constants
# =========================================================

GENERATED_DIR = "generated"


# =========================================================
# File / Filename Helpers
# =========================================================

def safe_filename(name: str) -> str:
    """
    Convert a candidate name into a safe filename.

    Example:
    'Jane Doe' -> 'jane_doe'
    """
    name = name.strip().lower()
    name = re.sub(r"[^a-z0-9]+", "_", name)
    return name.strip("_") or "cover_letter"


# =========================================================
# Main DOCX Export Entry Point
# =========================================================

def create_cover_letter_docx(full_name: str, cover_letter_text: str) -> str:
    """
    Create a DOCX cover letter file from plain-text cover letter content.

    Flow:
    1. Create output folder if needed
    2. Build Word document
    3. Apply default font styling
    4. Split plain text into paragraphs
    5. Save document to generated directory

    Expected cover letter structure:
    - line 1: candidate name
    - line 2: contact line
    - remaining lines: greeting, body paragraphs, closing
    """
    os.makedirs(GENERATED_DIR, exist_ok=True)

    filename = f"{safe_filename(full_name)}_cover_letter.docx"
    file_path = os.path.join(GENERATED_DIR, filename)

    document = Document()

    # =====================================================
    # Default Document Styling
    # =====================================================

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    # Remove blank lines and normalize paragraph input
    lines = [line.strip() for line in cover_letter_text.split("\n") if line.strip()]

    # Fallback structure if the cover letter text is unexpectedly empty
    if not lines:
        lines = [full_name, "Dear Hiring Manager,", "Sincerely,", full_name]

    # =====================================================
    # Header Rendering
    # =====================================================

    # First line = candidate name
    name_line = lines[0]

    # Second line usually = contact info
    contact_line = lines[1] if len(lines) > 1 else ""

    p = document.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    run = p.add_run(name_line)
    run.bold = True
    run.font.size = Pt(14)

    if contact_line:
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        run = p.add_run(contact_line)
        run.font.size = Pt(10)

    # =====================================================
    # Body Rendering
    # =====================================================

    remaining = lines[2:] if len(lines) > 2 else []

    for line in remaining:
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        run = p.add_run(line)
        run.font.size = Pt(11)

    document.save(file_path)
    return file_path